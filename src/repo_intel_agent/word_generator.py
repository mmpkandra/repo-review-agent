"""Generate ISMS-compliant Word documents from documentation artifacts."""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

from .models import DocumentationArtifact, RepositoryProfile


class ISMSDocumentGenerator:
    """Generate Word documents with ISMS-compliant formatting."""

    def __init__(
        self,
        organization: str = "Organization Name",
        department: str = "Engineering",
        classification: str = "Internal",
        author: str = "RunBook-AI",
    ) -> None:
        self.organization = organization
        self.department = department
        self.classification = classification
        self.author = author
        self.version = "1.0"
        self.date = datetime.now().strftime("%Y-%m-%d")

    def generate_document(
        self,
        artifact: DocumentationArtifact,
        profile: RepositoryProfile,
        output_path: Path,
    ) -> Path:
        """Generate a single Word document from an artifact."""
        doc = Document()
        self._setup_styles(doc)

        # Document ID based on artifact type
        doc_id = f"DOC-{profile.name.upper()[:10]}-{artifact.filename.replace('.md', '').replace('_', '-')[:15]}"

        self._add_cover_page(doc, artifact.title, doc_id, profile)
        self._add_document_control(doc, doc_id)
        self._add_version_history(doc)
        self._add_table_of_contents(doc)
        self._add_content(doc, artifact.body)

        # Save
        output_file = output_path / artifact.filename.replace(".md", ".docx")
        doc.save(str(output_file))
        return output_file

    def generate_all(
        self,
        artifacts: list[DocumentationArtifact],
        profile: RepositoryProfile,
        output_path: Path,
    ) -> list[Path]:
        """Generate Word documents for all artifacts."""
        output_path.mkdir(parents=True, exist_ok=True)
        generated = []
        for artifact in artifacts:
            path = self.generate_document(artifact, profile, output_path)
            generated.append(path)
        return generated

    def _setup_styles(self, doc: Document) -> None:
        """Configure document styles."""
        styles = doc.styles

        # Title style
        if "Doc Title" not in [s.name for s in styles]:
            title_style = styles.add_style("Doc Title", WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(28)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 51, 102)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)

    def _add_cover_page(
        self, doc: Document, title: str, doc_id: str, profile: RepositoryProfile
    ) -> None:
        """Add ISMS-compliant cover page."""
        # Organization header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run(self.organization)
        run.font.size = Pt(16)
        run.font.bold = True

        doc.add_paragraph()
        doc.add_paragraph()

        # Document title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)

        # Repository name
        repo_para = doc.add_paragraph()
        repo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        repo_run = repo_para.add_run(f"Repository: {profile.name}")
        repo_run.font.size = Pt(14)
        repo_run.font.italic = True

        doc.add_paragraph()
        doc.add_paragraph()

        # Metadata table
        table = doc.add_table(rows=6, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        metadata = [
            ("Document ID", doc_id),
            ("Version", self.version),
            ("Date", self.date),
            ("Classification", self.classification),
            ("Department", self.department),
            ("Author", self.author),
        ]

        for i, (label, value) in enumerate(metadata):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            # Bold labels
            row.cells[0].paragraphs[0].runs[0].font.bold = True

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(2)
            row.cells[1].width = Inches(3)

        doc.add_page_break()

    def _add_document_control(self, doc: Document, doc_id: str) -> None:
        """Add document control section."""
        doc.add_heading("Document Control", level=1)

        doc.add_heading("Purpose", level=2)
        doc.add_paragraph(
            "This document provides technical documentation for the repository. "
            "It is intended for engineering teams and stakeholders who need to "
            "understand, operate, or maintain the system."
        )

        doc.add_heading("Scope", level=2)
        doc.add_paragraph(
            "This document covers the technical aspects of the repository including "
            "architecture, configuration, deployment procedures, and operational guidelines."
        )

        doc.add_heading("Document Owner", level=2)
        doc.add_paragraph(f"Department: {self.department}")
        doc.add_paragraph(f"Classification: {self.classification}")

        doc.add_heading("Review Schedule", level=2)
        doc.add_paragraph("This document should be reviewed quarterly or when significant changes occur.")

        doc.add_page_break()

    def _add_version_history(self, doc: Document) -> None:
        """Add version history table."""
        doc.add_heading("Version History", level=1)

        table = doc.add_table(rows=2, cols=4)
        table.style = "Table Grid"

        # Header row
        headers = ["Version", "Date", "Author", "Changes"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        # First version entry
        row = table.rows[1]
        row.cells[0].text = self.version
        row.cells[1].text = self.date
        row.cells[2].text = self.author
        row.cells[3].text = "Initial document creation"

        doc.add_paragraph()
        doc.add_page_break()

    def _add_table_of_contents(self, doc: Document) -> None:
        """Add table of contents placeholder."""
        doc.add_heading("Table of Contents", level=1)
        doc.add_paragraph(
            "[Table of Contents - Update this field in Word after opening]"
        )
        doc.add_paragraph()
        doc.add_page_break()

    def _add_content(self, doc: Document, markdown_body: str) -> None:
        """Convert markdown content to Word format."""
        lines = markdown_body.split("\n")
        in_code_block = False
        in_table = False
        table_rows: list[list[str]] = []

        for line in lines:
            # Code blocks
            if line.startswith("```"):
                in_code_block = not in_code_block
                if not in_code_block and table_rows:
                    # End of code block
                    pass
                continue

            if in_code_block:
                para = doc.add_paragraph(line)
                para.style = "No Spacing"
                for run in para.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                continue

            # Tables
            if line.startswith("|"):
                if "---" in line:
                    continue  # Skip separator row
                cells = [c.strip() for c in line.split("|")[1:-1]]
                table_rows.append(cells)
                continue
            elif table_rows:
                # End of table, render it
                self._add_table(doc, table_rows)
                table_rows = []

            # Headings
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("#### "):
                doc.add_heading(line[5:], level=4)
            # Bullet points
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("  - ") or line.startswith("  * "):
                para = doc.add_paragraph(line[4:], style="List Bullet 2")
            # Numbered lists
            elif line and line[0].isdigit() and ". " in line[:4]:
                text = line.split(". ", 1)[1] if ". " in line else line
                doc.add_paragraph(text, style="List Number")
            # Regular paragraphs
            elif line.strip():
                para = doc.add_paragraph()
                self._add_formatted_text(para, line)
            else:
                # Empty line
                doc.add_paragraph()

        # Handle remaining table
        if table_rows:
            self._add_table(doc, table_rows)

    def _add_table(self, doc: Document, rows: list[list[str]]) -> None:
        """Add a table to the document."""
        if not rows:
            return

        num_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"

        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = table.rows[i].cells[j]
                    cell.text = cell_text.strip("`")
                    # Bold header row
                    if i == 0:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.bold = True

        doc.add_paragraph()

    def _add_formatted_text(self, paragraph, text: str) -> None:
        """Add text with basic formatting (bold, italic, code)."""
        # Simple parsing for **bold**, *italic*, `code`
        import re

        # Handle inline code
        parts = re.split(r'(`[^`]+`)', text)
        for part in parts:
            if part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            elif "**" in part:
                # Handle bold
                bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
                for bp in bold_parts:
                    if bp.startswith("**") and bp.endswith("**"):
                        run = paragraph.add_run(bp[2:-2])
                        run.font.bold = True
                    else:
                        paragraph.add_run(bp)
            else:
                paragraph.add_run(part)
